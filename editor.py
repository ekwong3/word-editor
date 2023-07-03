import sys
import os
import string
from docx import Document


def read_doc(doc):
    print('reading', doc)
    document = Document(doc)
    for section in document.sections:
        for p in section.first_page_header.paragraphs:
            print(p.text)
        for p in section.header.paragraphs:
            print(p.text)
        for p in section.even_page_header.paragraphs:
            print(p.text)

    for paragraph in document.paragraphs:
        print(paragraph.text)

    for section in document.sections:
        for p in section.first_page_footer.paragraphs:
            print(p.text)
        for p in section.footer.paragraphs:
            print(p.text)
        for p in section.even_page_footer.paragraphs:
            print(p.text)


def is_punc(char):
    return char in string.punctuation + string.whitespace


def matches_word(old_text, index, word_length):
    if (index == 0):
        return is_punc(old_text[word_length])
    elif (index + word_length == len(old_text)):
        return is_punc(old_text[index - 1])
    else:
        return is_punc(old_text[index - 1]) and is_punc(old_text[index + word_length])


def get_new_text(old_text, find_word, replace_word, match_case, match_word):
    word_length = len(find_word)
    if (match_case):
        if (match_word):
            res = ""
            ind = old_text.find(find_word)
            while ind >= 0:
                if matches_word(old_text, ind, word_length):
                    res += old_text[:ind]
                    res += replace_word
                else:
                    res += old_text[:ind + word_length]
                old_text = old_text[ind + word_length:]
                ind = old_text.find(find_word)
            res += old_text
            return res  # both match
        else:
            return old_text.replace(find_word, replace_word)  # only match case
    else:
        text_copy = old_text.lower()
        find_copy = find_word.lower()
        if (match_word):
            res = ""
            ind = text_copy.find(find_copy)
            while ind >= 0:
                if matches_word(text_copy, ind, word_length):
                    res += old_text[:ind]
                    res += replace_word
                else:
                    res += old_text[:ind + word_length]
                text_copy = text_copy[ind + word_length:]
                old_text = old_text[ind + word_length:]
                ind = text_copy.find(find_copy)
            res += old_text
            return res  # only match word
        else:
            res = ""
            i = 0
            while i < len(old_text):
                if old_text[i:i + word_length].lower() == find_copy:
                    res += replace_word
                    i = i + word_length
                else:
                    res += old_text[i]
                    i += 1
            return res  # don't have to match case or word


def find_and_replace(doc, find, replace, match_case, match_word):
    document = Document(doc)
    sections = document.sections
    for i in range(len(document.paragraphs)):
        paragraph_text = document.paragraphs[i].text
        new_text = get_new_text(paragraph_text, find,
                                replace, match_case, match_word)
        document.paragraphs[i].text = new_text

    for section in sections:
        if section.different_first_page_header_footer:
            diff_header = section.first_page_header
            for i in range(len(diff_header.paragraphs)):
                header_text = diff_header.paragraphs[i].text
                new_text = get_new_text(header_text, find,
                                        replace, match_case, match_word)
                diff_header.paragraphs[i].text = new_text

            diff_footer = section.first_page_footer
            for i in range(len(diff_footer.paragraphs)):
                footer_text = diff_footer.paragraphs[i].text
                new_text = get_new_text(footer_text, find,
                                        replace, match_case, match_word)
                diff_footer.paragraphs[i].text = new_text

        if document.settings.odd_and_even_pages_header_footer:
            diff_header = section.even_page_header
            for i in range(len(diff_header.paragraphs)):
                header_text = diff_header.paragraphs[i].text
                new_text = get_new_text(header_text, find,
                                        replace, match_case, match_word)
                diff_header.paragraphs[i].text = new_text

            diff_footer = section.even_page_footer
            for i in range(len(diff_footer.paragraphs)):
                footer_text = diff_footer.paragraphs[i].text
                new_text = get_new_text(footer_text, find,
                                        replace, match_case, match_word)
                diff_footer.paragraphs[i].text = new_text

        header = section.header
        for i in range(len(header.paragraphs)):
            header_text = header.paragraphs[i].text
            new_text = get_new_text(header_text, find,
                                    replace, match_case, match_word)
            header.paragraphs[i].text = new_text

        footer = section.footer
        for i in range(len(footer.paragraphs)):
            footer_text = footer.paragraphs[i].text
            new_text = get_new_text(footer_text, find,
                                    replace, match_case, match_word)
            footer.paragraphs[i].text = new_text

    document.save(doc)


def replace_many(doc, finds, replaces, match_case, match_word):
    finds = finds.split('\n')
    replaces = replaces.split('\n')
    shorter = min(len(finds), len(replaces))
    for i in range(shorter):
        find_and_replace(doc, finds[i], replaces[i], match_case, match_word)


def replace_folder(folderpath, finds, replaces, match_case, match_word):
    for filename in os.listdir(folderpath):
        if filename.endswith('.docx'):
            filepath = f'{folderpath}{filename}'
            replace_many(filepath, finds, replaces, match_case, match_word)


folder = sys.argv[1]
find = sys.argv[2]
replace = sys.argv[3]
match_case = sys.argv[4] == "true"
match_word = sys.argv[5] == "true"

replace_folder(folder, find, replace, match_case, match_word)
