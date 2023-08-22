import sys
import os
import string
from docx import Document

NON_PUNCTUATION = string.ascii_letters + string.digits


def read_headers(document):
    print('reading headers')
    for section in document.sections:
        for p in section.first_page_header.paragraphs:
            print(p.text)
        for p in section.header.paragraphs:
            print(p.text)
        for p in section.even_page_header.paragraphs:
            print(p.text)


def read_paragraphs(document):
    print('reading paragraphs')
    for paragraph in document.paragraphs:
        print(paragraph.text)


def read_tables(document):
    print('reading tables')
    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                print(table.cell(row, col).text)


def read_footers(document):
    print('reading footers')
    for section in document.sections:
        for p in section.first_page_footer.paragraphs:
            print(p.text)
        for p in section.footer.paragraphs:
            print(p.text)
        for p in section.even_page_footer.paragraphs:
            print(p.text)


def read_doc(doc):
    print('reading', doc)
    document = Document(doc)
    read_headers(document)
    read_paragraphs(document)
    read_tables(document)
    read_footers(document)


def is_punc(char):
    return char.lower() not in NON_PUNCTUATION


def matches_word(old_text, index, word_length):
    if (index + word_length == len(old_text) and index == 0):
        return True
    elif (index + word_length == len(old_text)):
        return is_punc(old_text[index - 1])
    elif (index == 0):
        return is_punc(old_text[word_length])
    else:
        return is_punc(old_text[index - 1]) and is_punc(old_text[index + word_length])


def get_new_word(old_word, replace_word):
    if (old_word.islower()):
        return replace_word.lower()
    elif (old_word.isupper()):
        return replace_word.upper()
    elif (old_word[0].isupper() and old_word[1:].islower()):
        return replace_word[0].upper() + replace_word[1:].lower()
    else:
        return replace_word


def replace_single(run, start_index, find_length, replace_word, keep_case, match_word):
    if (match_word and matches_word(run.text, start_index, find_length)) or not match_word:
        if keep_case:
            replace_word = get_new_word(
                run.text[start_index: start_index + find_length + 1], replace_word)
        run.text = run.text[: start_index] + replace_word + \
            run.text[start_index + find_length:]


def replace_multiple(paragraph, run_index, start_index, find_length, replace_word,
                     keep_case, match_word):
    if run_index + 1 == len(paragraph.runs):
        replace_single(paragraph.runs[run_index], start_index,
                       find_length, replace_word, keep_case, match_word)
        return
    current_run = paragraph.runs[run_index]
    combined_text = current_run.text
    final_length = start_index + find_length
    next_index = run_index + 1
    while len(combined_text) <= final_length and next_index < len(paragraph.runs):
        combined_text += paragraph.runs[next_index].text
        next_index += 1

    if (match_word and matches_word(combined_text, start_index, find_length)) or not match_word:
        if keep_case:
            replace_word = get_new_word(
                combined_text[start_index: start_index + find_length + 1], replace_word)

        remaining_length = find_length - (len(current_run.text) - start_index)
        next_index = run_index + 1
        while remaining_length > 0:
            next_run = paragraph.runs[next_index]
            run_length = len(next_run.text)
            if 0 < run_length < remaining_length:
                remaining_length -= run_length
                next_run.text = ''
            elif run_length >= remaining_length:
                next_run.text = next_run.text[remaining_length:]
                remaining_length = 0
            next_index += 1
        current_run.text = current_run.text[:start_index] + replace_word


def replace_text(paragraph, find_word, replace_word, keep_case, match_word):
    if len(paragraph.runs) < 1:
        return
    word_length = len(find_word)
    text_copy = paragraph.text.lower()
    find_copy = find_word.lower()
    word_start = text_copy.find(find_copy)
    text_index = 0
    run_index = 0
    current_run_length = len(paragraph.runs[run_index].text)
    while word_start >= 0:
        while text_index + current_run_length <= word_start:
            run_index += 1
            text_index += current_run_length
            current_run_length = len(paragraph.runs[run_index].text)
        run_start = word_start - text_index
        # Want to be strictly less than to account for match_word
        if (run_start + word_length < len(paragraph.runs[run_index].text)):
            replace_single(paragraph.runs[run_index], run_start,
                           word_length, replace_word, keep_case, match_word)
        else:
            replace_multiple(paragraph, run_index, run_start,
                             word_length, replace_word, keep_case, match_word)
        text_copy = paragraph.text.lower()
        word_start = text_copy.find(find_copy)
        current_run_length = len(paragraph.runs[run_index].text)


def find_and_replace_section(section, find, replace, keep_case, match_word):
    for i in range(len(section.paragraphs)):
        replace_text(section.paragraphs[i], find,
                     replace, keep_case, match_word)


def find_and_replace(doc, find, replace, keep_case, match_word):
    document = Document(doc)
    sections = document.sections
    find_and_replace_section(document, find, replace, keep_case, match_word)

    for section in sections:
        header = section.header
        find_and_replace_section(header, find, replace, keep_case, match_word)

        footer = section.footer
        find_and_replace_section(
            footer, find, replace, keep_case, match_word)

        if section.different_first_page_header_footer:
            diff_header = section.first_page_header
            find_and_replace_section(
                diff_header, find, replace, keep_case, match_word)

            diff_footer = section.first_page_footer
            find_and_replace_section(
                diff_footer, find, replace, keep_case, match_word)

        if document.settings.odd_and_even_pages_header_footer:
            diff_header = section.even_page_header
            find_and_replace_section(
                diff_header, find, replace, keep_case, match_word)

            diff_footer = section.even_page_footer
            find_and_replace_section(
                diff_footer, find, replace, keep_case, match_word)

    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                find_and_replace_section(table.cell(
                    row, col), find, replace, keep_case, match_word)

    document.save(doc)


def replace_many(doc, finds, replaces, keep_case, match_word):
    finds = finds.split('\n')
    replaces = replaces.split('\n')
    shorter = min(len(finds), len(replaces))
    for i in range(shorter):
        find_and_replace(doc, finds[i], replaces[i], keep_case, match_word)


def replace_folder(folderpath, finds, replaces, keep_case, match_word, process_sub):
    for path in os.listdir(folderpath):
        # omit non word files and temporary open ones
        if path.endswith('.docx') and path[:2] != '~$':
            filepath = f'{folderpath}{path}'
            replace_many(filepath, finds, replaces, keep_case, match_word)
        elif process_sub and os.path.isdir(f'{folderpath}{path}/'):
            replace_folder(f'{folderpath}{path}/', finds,
                           replaces, keep_case, match_word, process_sub)


doc = '/Users/ekwong/Desktop/test.docx'
d = Document(doc)

# folder = sys.argv[1]
# find = sys.argv[2]
# replace = sys.argv[3]
# keep_case = sys.argv[4] == "true"
# match_word = sys.argv[5] == "true"
# process_sub = sys.argv[6] == "true"

# replace_folder(folder, find, replace, keep_case, match_word, process_sub)
